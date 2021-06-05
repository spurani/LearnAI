# -*- coding: utf-8 -*-
"""
Created on Sun Jan  3 11:03:53 2021

@author: SP
"""

import os
import docx
from docx import Document
document = Document()
import pandas as pd
os.getcwd()
#os.mkdir("Tester")
os.chdir("Tester")
print(os.getcwd())

textfile = open("Greetings.txt","w+")
textfile.write("Welcome To Text Mining and Natural Lanuguae Processing")
textfile.close()

os.rename("C:/Users/SP.000/Documents/Edureka/Machine Learning Engineer/Natural Language Processing/Module 1 Introduction to Text Mining and NLP/Tester/Greetings.txt","C:/Users/SP.000/Documents/Edureka/Machine Learning Engineer/Natural Language Processing/Module 1 Introduction to Text Mining and NLP/Tester/Welcome.txt")

doc = docx.Document('C:/Users/SP.000/Documents/Edureka/Machine Learning Engineer/Natural Language Processing/Module 1 Introduction to Text Mining and NLP/698_m1_datasets_v1.0/NLP.docx')
print(len(doc.paragraphs))
#for a in doc.paragraphs:
#    print(a.runs[0])
txt = open("Welcome.txt","r")
#print(txt)
paragraph = doc.add_paragraph("lodu")
prior_paragraph = paragraph.insert_paragraph_before('To the outside observer')
doc.save('C:/Users/SP.000/Documents/Edureka/Machine Learning Engineer/Natural Language Processing/Module 1 Introduction to Text Mining and NLP/698_m1_datasets_v1.0/NLP.docx')
print(doc.paragraphs[0].runs[0].text)

file = pd.read_csv('C:/Users/SP.000/Documents/Edureka/Machine Learning Engineer/Natural Language Processing/Module 1 Introduction to Text Mining and NLP/698_m1_datasets_v1.0/EmployeeDetails.csv')
file[['First Name','Last Name']] = file['N'].str.split(" ",expand=True)
print(file)
del(file['N'])
print(file)
file['Salary'] = (file.loc[:,'Salary']) * (10/100) + ((file.loc[:,'Salary']))
print(file)
file.to_csv("C:/Users/SP.000/Documents/Edureka/Machine Learning Engineer/Natural Language Processing/Module 1 Introduction to Text Mining and NLP/698_m1_datasets_v1.0/Employee_Data.csv",index=False)